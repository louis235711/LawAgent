import os
import re
from collections.abc import AsyncGenerator
from typing import Union

from docx import Document as DocxDocument
from src.agents.base import BaseAgent, AgentResponse
from src.llm.client import chat_completion
from src.config import settings
from loguru import logger

WRITING_PROMPT = """你是一个法律文书撰写专家。请根据用户的需求，撰写一份完整的法律文书。

## 要求
1. 第一行必须是文书的标题，以单个 # 开头（例如：# 借款合同），标题需简短概括文书内容
2. 使用标准的法律文书格式和法言法语
3. 结构完整，包含标题、正文、落款等必要部分
4. 如果用户未提供某些信息，用【待补充：字段名】标记
5. 在文书末尾列出需要用户确认或补充的字段清单
6. 不得编造不存在的法律条款

## 用户需求
{user_input}

## 完整文书
"""


class DocumentWritingAgent(BaseAgent):
    def __init__(self):
        super().__init__("document_writing")

    async def execute(
        self,
        session_id: str,
        user_input: str,
        context: list[dict],
    ) -> AgentResponse:
        logger.info(f"[AGENT] document_writing session={session_id}")
        return await self._generate(session_id, user_input, context)

    async def stream_execute(
        self, session_id: str, user_input: str, context: list[dict],
    ) -> AsyncGenerator[Union[str, AgentResponse], None]:
        logger.info(f"[AGENT] document_writing stream session={session_id}")
        result = await self._generate(session_id, user_input, context)
        yield result

    async def _generate(
        self, session_id: str, user_input: str, context: list[dict],
    ) -> AgentResponse:
        # 1. Determine format
        fmt = "docx"
        if "md" in user_input.lower() or "markdown" in user_input.lower():
            fmt = "md"
        elif "txt" in user_input.lower() or "纯文本" in user_input.lower() or "文本文件" in user_input.lower():
            fmt = "txt"

        # 2. Generate document via LLM
        prompt = WRITING_PROMPT.format(user_input=user_input)
        messages = context + [{"role": "user", "content": prompt}]

        content = await chat_completion(messages=messages, temperature=0.5, max_tokens=4096)
        logger.info(f"[AGENT] document_writing generated: {len(content)} chars")

        # 3. Extract title from the first # heading line
        title = _extract_title(content)
        safe_title = re.sub(r'[\\/*?:"<>|]', '', title).strip()[:50]
        if not safe_title:
            safe_title = "法律文书"

        filename = f"{safe_title}.{fmt}"
        output_dir = os.path.join(settings.generated_dir, session_id)
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, filename)

        # 4. Export to disk
        if fmt == "docx":
            export_to_docx(content, output_path)
        else:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)

        logger.info(f"[AGENT] document_writing saved: {output_path}")
        download_url = f"/api/download/{session_id}/{filename}"

        return AgentResponse(
            content=f"文档已生成：**{title}**",
            memory_content=content,
            metadata={
                "message_type": "文书",
                "download_url": download_url,
                "filename": filename,
                "title": title,
                "format": fmt,
            },
        )


def _extract_title(content: str) -> str:
    for line in content.split("\n"):
        stripped = line.strip()
        m = re.match(r'^#+\s*(.+)', stripped)
        if m:
            return m.group(1).strip()
    # Fallback: first non-empty line
    for line in content.split("\n"):
        stripped = line.strip()
        if stripped:
            return stripped[:50]
    return "法律文书"


def export_to_docx(content: str, output_path: str):
    doc = DocxDocument()
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        else:
            doc.add_paragraph(stripped)
    doc.save(output_path)
