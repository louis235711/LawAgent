"""Tool registry for ReAct agent — Anthropic tool_use format."""

import os
import re
import json
import math
from datetime import datetime
from src.rag.pipeline import retrieve_legal, retrieve_session_docs
from src.tools.web_search import search_web
from src.config import settings
from loguru import logger

# ── Tool definitions (Anthropic format) ──────────────────────

TOOLS = [
    {
        "name": "search_laws",
        "description": "搜索中国法律知识库，获取相关法条原文、章节和条款号。调用时机：用户咨询具体法律问题时必须先检索再回答。",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "法律检索查询，使用法律专业术语，如'竞业限制 违约金'而非口语化表达",
                },
                "top_k": {
                    "type": "integer",
                    "description": "返回结果数量，默认5条",
                    "default": 5,
                },
            },
            "required": ["query"],
        },
    },
    {
        "name": "search_cases",
        "description": "联网搜索类似案例的判决结果和处理方式。调用时机：案情分析需要类案参考、了解实务倾向时。",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "案例搜索查询，如'民间借贷 利息 判决 2024'",
                },
                "max_results": {
                    "type": "integer",
                    "description": "最大结果数，默认5条",
                    "default": 5,
                },
            },
            "required": ["query"],
        },
    },
    {
        "name": "search_documents",
        "description": "在用户已上传的文档中搜索相关内容（向量+关键词混合检索）。调用时机：用户针对已上传文档的具体内容提问时。",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "文档搜索查询，提取问题中的关键概念",
                },
                "top_k": {
                    "type": "integer",
                    "description": "返回结果数量，默认5条",
                    "default": 5,
                },
            },
            "required": ["query"],
        },
    },
    {
        "name": "read_document_full",
        "description": "读取用户已上传文档的全文。调用时机：合同审查、需要完整阅读整份文档而非片段检索时。",
        "input_schema": {
            "type": "object",
            "properties": {},
            "required": [],
        },
    },
    {
        "name": "generate_document",
        "description": "根据用户需求生成完整的法律文书（合同、起诉状、律师函、申请书等）。调用时机：用户明确要求起草/撰写/生成法律文书。",
        "input_schema": {
            "type": "object",
            "properties": {
                "requirements": {
                    "type": "string",
                    "description": "文书需求描述，包含文书类型、当事人信息、关键条款等，越详细越好",
                },
                "format": {
                    "type": "string",
                    "enum": ["md", "docx", "pdf", "txt"],
                    "description": "输出格式，默认md。pdf 输出渲染后的 PDF 文件，排版美观适合打印。",
                },
            },
            "required": ["requirements"],
        },
    },
    {
        "name": "get_current_time",
        "description": "获取当前日期和时间。调用时机：用户询问'今天是什么日期'、计算时间差、诉讼时效计算等需要准确日期的场景。",
        "input_schema": {
            "type": "object",
            "properties": {},
            "required": [],
        },
    },
    {
        "name": "calculate",
        "description": "执行数学计算。调用时机：需要计算利息、违约金、赔偿金额、诉讼费用等涉及数学运算的场景。支持四则运算、百分比、幂运算等。",
        "input_schema": {
            "type": "object",
            "properties": {
                "expression": {
                    "type": "string",
                    "description": "数学表达式，如 '(100000 * 0.05 * 365) / 365' 或 '50000 * 1.2 + 3000'",
                },
            },
            "required": ["expression"],
        },
    },
    {
        "name": "search_memory",
        "description": "检索用户的长期记忆（跨会话的历史对话摘要）。调用时机：用户提到之前讨论过的内容、需要回顾历史对话结论、或当前问题需要参考过去的咨询记录时。仅返回7天内的记忆。",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "记忆检索查询，如'之前讨论的借款合同''上次咨询的劳动纠纷结论'",
                },
                "top_k": {
                    "type": "integer",
                    "description": "返回结果数量，默认5条",
                    "default": 5,
                },
            },
            "required": ["query"],
        },
    },
]

# ── PDF HTML template (xhtml2pdf-compatible legal styling) ───

WKHTMLTOPDF_PATH = r"C:\Users\lenovo\AppData\Local\wkhtmltopdf\wkhtmltopdf.exe"

PDF_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<style>
  @page {{ size: A4; margin: 2.5cm 2cm 2.5cm 2cm; }}
  @font-face {{ font-family: 'SimHei'; src: url('file:///{fonts_dir}/simhei.ttf') format('truetype'); }}
  @font-face {{ font-family: 'SimFang'; src: url('file:///{fonts_dir}/simfang.ttf') format('truetype'); }}
  body {{ font-family: 'SimFang', 'SimHei', serif; font-size: 12pt; line-height: 1.8; color: #222; }}
  h1 {{ font-family: 'SimHei', sans-serif; font-size: 18pt; text-align: center; margin: 0 0 1em 0; }}
  h2 {{ font-family: 'SimHei', sans-serif; font-size: 14pt; margin: 1.2em 0 0.5em 0; border-bottom: 1px solid #ccc; padding-bottom: 0.2em; }}
  h3 {{ font-family: 'SimHei', sans-serif; font-size: 12.5pt; margin: 1em 0 0.4em 0; }}
  p {{ margin: 0.5em 0; text-indent: 2em; }}
  ul, ol {{ margin: 0.5em 0 0.5em 1em; padding-left: 1em; }}
  li {{ margin: 0.2em 0; }}
  table {{ border-collapse: collapse; width: 100%; margin: 0.8em 0; font-size: 10.5pt; }}
  th, td {{ border: 1px solid #555; padding: 6px 8px; text-align: left; }}
  th {{ background: #e8e8e8; font-weight: bold; }}
  pre {{ background: #f4f4f4; border: 1px solid #ddd; border-left: 3px solid #4361ee; padding: 10px 14px; font-family: Consolas, monospace; font-size: 9pt; line-height: 1.5; }}
  code {{ font-family: Consolas, monospace; font-size: 9.5pt; background: #f0f0f0; padding: 1px 4px; }}
  pre code {{ background: none; padding: 0; }}
  blockquote {{ border-left: 3px solid #4361ee; margin: 0.8em 0; padding: 0.5em 1em; background: #f8f9ff; color: #444; }}
  strong {{ font-weight: bold; }}
  em {{ font-style: italic; }}
  hr {{ border: none; border-top: 1px solid #ddd; margin: 1.5em 0; }}
</style>
</head>
<body>
{content}
</body>
</html>"""


# ── Tool executors ───────────────────────────────────────────

async def execute_search_laws(query: str, top_k: int = 5, history_text: str = "") -> str:
    """Execute legal knowledge search. Returns formatted law articles."""
    try:
        results = await retrieve_legal(query, top_k=top_k, history_text=history_text)
    except Exception as e:
        logger.warning(f"search_laws failed: {e}")
        return f"法条检索失败: {e}"

    if not results:
        return "未找到相关法条。法律知识库可能尚未收录相关内容。"

    lines = []
    for i, r in enumerate(results, 1):
        law_name = r.get("law_name", "")
        article = r.get("article_number", "")
        chapter = r.get("chapter", "")
        header = f"《{law_name}》" if law_name else ""
        if article:
            header += f" 第{article}条"
        if chapter:
            header += f"（{chapter}）"
        lines.append(f"### 法条{i}: {header}\n{r['chunk_text']}")

    return "\n\n".join(lines)


async def execute_search_cases(query: str, max_results: int = 5) -> str:
    """Execute web search for similar cases. Returns formatted case references."""
    try:
        results = await search_web(query, max_results=max_results)
    except Exception as e:
        logger.warning(f"search_cases failed: {e}")
        return f"类案搜索失败: {e}"

    if not results:
        return "未找到类案参考。"

    lines = []
    for i, r in enumerate(results, 1):
        lines.append(f"### 案例{i}: {r['title']}\n{r['content']}\n来源: {r['url']}")

    return "\n\n".join(lines)


async def execute_search_documents(query: str, session_id: str, top_k: int = 5, history_text: str = "") -> str:
    """Execute session document search. Returns formatted document excerpts."""
    try:
        results = await retrieve_session_docs(query, session_id, top_k=top_k, history_text=history_text)
    except Exception as e:
        logger.warning(f"search_documents failed: {e}")
        return f"文档检索失败: {e}"

    if not results:
        return "未在文档中找到相关内容。"

    lines = []
    for i, r in enumerate(results, 1):
        lines.append(f"### 文档相关内容{i}\n{r['chunk_text']}")

    return "\n\n---\n\n".join(lines)


async def execute_read_document_full(session_id: str, user_id: int = 0) -> str:
    """Read full document text for a session. Falls back to JSON chunks if not in Milvus."""
    try:
        from src.vector_db.milvus_client import get_collection, SESSION_DOCUMENTS_COLLECTION

        coll = get_collection(SESSION_DOCUMENTS_COLLECTION)
        offset = 0
        batch = 500
        rows = []
        while True:
            chunk_results = coll.query(
                expr=f'session_id == "{session_id}"',
                output_fields=["chunk_text", "chunk_index"],
                limit=batch,
                offset=offset,
            )
            if not chunk_results:
                break
            rows.extend(chunk_results)
            if len(chunk_results) < batch:
                break
            offset += batch

        if rows:
            rows.sort(key=lambda r: r.get("chunk_index", 0))
            total = len(rows)
            # If too many chunks, return overview
            if total > 50:
                summary_parts = []
                for r in rows[:50]:
                    idx = r.get("chunk_index", 0)
                    text = r.get("chunk_text") or ""
                    summary_parts.append(f"【第{idx+1}段】{text}")
                return (
                    f"文档共 {total} 段（以下为前 50 段摘要，后续可在后续轮次按需检索）:\n\n"
                    + "\n\n---\n\n".join(summary_parts)
                )
            parts = []
            for r in rows:
                idx = r.get("chunk_index", 0)
                text = r.get("chunk_text") or ""
                parts.append(f"【第{idx+1}段】{text}")
            return "\n\n---\n\n".join(parts)

        # Fallback: load from JSON (check user-scoped first, then legacy)
        import glob

        def _find_session_dir(base_dir: str) -> str | None:
            if user_id:
                user_dir = os.path.join(base_dir, str(user_id), session_id)
                if os.path.isdir(user_dir):
                    return user_dir
            legacy_dir = os.path.join(base_dir, session_id)
            if os.path.isdir(legacy_dir):
                return legacy_dir
            return None

        session_dir = _find_session_dir(settings.uploads_dir)
        if session_dir:
            json_files = glob.glob(os.path.join(session_dir, "*_chunks.json"))
            if json_files:
                with open(json_files[0], "r", encoding="utf-8") as f:
                    chunks = json.load(f)
                total = len(chunks)
                if total > 50:
                    return f"文档共 {total} 段（前50段）:\n\n" + "\n\n---\n\n".join(chunks[:50])
                return "\n\n---\n\n".join(chunks)

            txt_files = glob.glob(os.path.join(session_dir, "*.txt"))
            if txt_files:
                with open(txt_files[0], "r", encoding="utf-8") as f:
                    return f.read()

        return "文档未找到，可能已过期或未正确上传。"
    except Exception as e:
        logger.error(f"read_document_full failed: {e}")
        return f"读取文档失败: {e}"

def _add_formatted_paragraph(doc, text: str):
    """Add a paragraph with inline markdown formatting (bold, italic, code)."""
    from docx.shared import Pt
    para = doc.add_paragraph()
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*|'
        r'\*\*(.+?)\*\*|'
        r'\*(.+?)\*|'
        r'~~(.+?)~~|'
        r'`(.+?)`|'
        r'__(.+?)__)'
    )
    last_end = 0
    for m in pattern.finditer(text):
        plain = text[last_end:m.start()]
        if plain:
            para.add_run(plain)
        groups = m.groups()
        if groups[0] is not None:
            run = para.add_run(groups[1])
            run.bold = True
            run.italic = True
        elif groups[2] is not None:
            para.add_run(groups[2]).bold = True
        elif groups[3] is not None:
            para.add_run(groups[3]).italic = True
        elif groups[4] is not None:
            run = para.add_run(groups[4])
            run.font.strike = True
        elif groups[5] is not None:
            run = para.add_run(groups[5])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            from docx.oxml.ns import qn
            from lxml import etree
            rPr = run._element.get_or_add_rPr()
            shd = etree.SubElement(rPr, qn('w:shd'))
            shd.set(qn('w:fill'), 'E8E8E8')
            shd.set(qn('w:val'), 'clear')
        elif groups[6] is not None:
            para.add_run(groups[6]).bold = True
        last_end = m.end()
    plain = text[last_end:]
    if plain:
        para.add_run(plain)


async def execute_generate_document(requirements: str, format: str, session_id: str, user_id: int = 0) -> str:
    """Generate a legal document. Returns markdown content + download info."""
    from src.llm.client import chat_completion

    WRITING_PROMPT = """你是法律文书撰写专家。请根据用户需求撰写完整法律文书。

## 要求
1. 第一行必须是文书标题，以单个 # 开头
2. 使用标准法律文书格式和法言法语
3. 结构完整（标题、正文、落款等）
4. 未提供的信息用【待补充：字段名】标记
5. 末尾列出需用户确认或补充的字段清单
6. 不得编造不存在的法律条款

## 用户需求
{requirements}

## 完整文书
"""
    try:
        content = await chat_completion(
            messages=[{"role": "user", "content": WRITING_PROMPT.format(requirements=requirements)}],
            temperature=0.5,
            max_tokens=4096,
        )
    except Exception as e:
        logger.error(f"generate_document LLM failed: {e}")
        return f"文书生成失败: {e}"

    # Extract title
    title = "法律文书"
    for line in content.split("\n"):
        m = re.match(r'^#+\s*(.+)', line.strip())
        if m:
            title = re.sub(r'[\\/*?:"<>|]', '', m.group(1).strip())[:50]
            break

    if not format:
        format = "md"

    filename = f"{title}.{format}"
    output_dir = os.path.join(settings.generated_dir, str(user_id), session_id) if user_id else os.path.join(settings.generated_dir, session_id)
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, filename)

    if format == "docx":
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches
        doc = DocxDocument()
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif re.match(r'^[-*]\s+', stripped):
                doc.add_paragraph(
                    re.sub(r'^[-*]\s+', '', stripped),
                    style='List Bullet'
                )
            elif re.match(r'^\d+[.、]\s*', stripped):
                doc.add_paragraph(
                    re.sub(r'^\d+[.、]\s*', '', stripped),
                    style='List Number'
                )
            else:
                _add_formatted_paragraph(doc, stripped)
        doc.save(output_path)
    elif format == "pdf":
        import markdown
        import tempfile
        import shutil
        import pdfkit
        # Copy fonts to temp dir (avoid non-ASCII path issues)
        fonts_dir = os.path.join(tempfile.gettempdir(), "lawagent_fonts")
        os.makedirs(fonts_dir, exist_ok=True)
        for fn, src in [("simhei.ttf", "C:/Windows/Fonts/simhei.ttf"),
                         ("simfang.ttf", "C:/Windows/Fonts/simfang.ttf")]:
            dst = os.path.join(fonts_dir, fn)
            if not os.path.exists(dst) and os.path.exists(src):
                shutil.copy2(src, dst)
        md_html = markdown.markdown(
            content, extensions=['tables', 'fenced_code', 'codehilite'],
        )
        styled = PDF_HTML_TEMPLATE.format(fonts_dir=fonts_dir.replace(os.sep, "/"), content=md_html)
        # Write HTML to temp file for wkhtmltopdf
        html_path = output_path + ".html"
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(styled)
        try:
            config = pdfkit.configuration(wkhtmltopdf=WKHTMLTOPDF_PATH)
            pdfkit.from_file(html_path, output_path, configuration=config,
                             options={'encoding': 'UTF-8', 'enable-local-file-access': ''})
        finally:
            if os.path.exists(html_path):
                os.remove(html_path)
    else:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)

    from urllib.parse import quote
    download_url = f"/api/download/{session_id}/{quote(filename)}"
    logger.info(f"generate_document: {title}.{format} saved")

    display_text = (
        f"## 文书已生成\n\n"
        f"**标题**: {title}\n"
        f"**格式**: {format}\n"
        f"**下载链接**: {download_url}\n\n"
        f"---\n\n"
        f"{content}"
    )
    return json.dumps({
        "display_text": display_text,
        "download_url": download_url,
        "filename": filename,
        "title": title,
        "format": format,
    }, ensure_ascii=False)


async def execute_get_current_time() -> str:
    """Return current date and time with weekday in Chinese."""
    now = datetime.now()
    weekdays = ["一", "二", "三", "四", "五", "六", "日"]
    wd = weekdays[now.weekday()]
    return (
        f"当前日期时间：{now.strftime('%Y年%m月%d日')} 星期{wd} {now.strftime('%H:%M:%S')}\n"
        f"ISO 格式：{now.isoformat()}\n"
        f"Unix 时间戳：{int(now.timestamp())}"
    )


# Safe math functions for calculate
_SAFE_MATH = {
    "abs": abs, "round": round, "min": min, "max": max,
    "int": int, "float": float,
    "sqrt": math.sqrt, "pow": pow,
    "ceil": math.ceil, "floor": math.floor,
    "sin": math.sin, "cos": math.cos, "tan": math.tan,
    "log": math.log, "log10": math.log10,
    "pi": math.pi, "e": math.e,
}


async def execute_calculate(expression: str) -> str:
    """Safely evaluate a mathematical expression."""
    try:
        result = eval(expression, {"__builtins__": {}}, _SAFE_MATH)
        return (
            f"表达式：{expression}\n"
            f"计算结果：{result}\n"
            f"（使用 Python 标准运算规则，精度为浮点数）"
        )
    except Exception as e:
        logger.warning(f"calculate failed: {e}, expr: {expression}")
        return f"计算失败：{e}。请检查表达式格式是否正确。"


async def execute_search_memory(query: str, user_id: int, top_k: int = 5) -> str:
    """Search user's long-term memory (structured summaries from past sessions)."""
    from src.rag.pipeline import retrieve_session_memory
    try:
        results = await retrieve_session_memory(query=query, user_id=user_id, top_k=top_k)
    except Exception as e:
        logger.warning(f"search_memory failed: {e}")
        return f"记忆检索失败: {e}"

    if not results:
        return "未找到相关历史记忆。可能已超过7天有效期或尚未积累足够记忆。"

    lines = []
    for i, r in enumerate(results, 1):
        topic = r.get("topic", "")
        header = f"### 记忆{i}"
        if topic:
            header += f"（{topic}）"
        lines.append(f"{header}\n{r['chunk_text']}")

    return "\n\n".join(lines)


# ── Tool executor router ─────────────────────────────────────

async def execute_tool(name: str, input_: dict, session_id: str, user_id: int = 0) -> str:
    """Route tool name to executor. Returns serialized result string."""
    # External MCP tools (prefixed with mcp_)
    if name.startswith("mcp_"):
        from src.mcp.client import get_mcp_client
        return await get_mcp_client().call_tool(name, input_)

    try:
        if name == "search_laws":
            return await execute_search_laws(
                query=input_.get("query", ""),
                top_k=input_.get("top_k", 5),
            )
        elif name == "search_cases":
            return await execute_search_cases(
                query=input_.get("query", ""),
                max_results=input_.get("max_results", 5),
            )
        elif name == "search_documents":
            return await execute_search_documents(
                query=input_.get("query", ""),
                session_id=session_id,
                top_k=input_.get("top_k", 5),
            )
        elif name == "read_document_full":
            return await execute_read_document_full(session_id, user_id=user_id)
        elif name == "generate_document":
            return await execute_generate_document(
                requirements=input_.get("requirements", ""),
                format=input_.get("format", "md"),
                session_id=session_id,
                user_id=user_id,
            )
        elif name == "get_current_time":
            return await execute_get_current_time()
        elif name == "calculate":
            return await execute_calculate(
                expression=input_.get("expression", ""),
            )
        elif name == "search_memory":
            return await execute_search_memory(
                query=input_.get("query", ""),
                user_id=user_id,
                top_k=input_.get("top_k", 5),
            )
        else:
            return f"未知工具: {name}"
    except Exception as e:
        logger.error(f"Tool execution error ({name}): {e}")
        return f"工具执行出错 ({name}): {e}"


def make_tool_result(tool_use_id: str, content: str) -> dict:
    """Build an Anthropic tool_result content block."""
    return {
        "type": "tool_result",
        "tool_use_id": tool_use_id,
        "content": content,
    }
