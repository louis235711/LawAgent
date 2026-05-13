"""Multi-format document parser: PDF/PNG/JPG via OCR, DOCX/XLSX via python-docx/openpyxl, MD/TXT direct read."""

import os
import httpx
import docx
import openpyxl
from src.config import settings
from loguru import logger


async def parse_document(file_path: str, filename: str) -> str:
    """Parse any supported document to plain text. Raises on error."""
    ext = os.path.splitext(filename)[1].lower()

    if ext in (".pdf", ".png", ".jpg", ".jpeg"):
        return await _parse_via_ocr(file_path, filename)
    elif ext == ".docx":
        return _parse_docx(file_path)
    elif ext == ".xlsx":
        return _parse_xlsx(file_path)
    elif ext in (".md", ".txt"):
        return _parse_text(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


async def _parse_via_ocr(file_path: str, filename: str) -> str:
    """Send to PaddleOCR Docker service."""
    url = f"{settings.ocr_service_url}/ocr"
    logger.info(f"OCR parsing: {filename} → {url}")

    try:
        async with httpx.AsyncClient(timeout=300) as client:
            with open(file_path, "rb") as f:
                response = await client.post(
                    url,
                    files={"file": (filename, f)},
                )
    except httpx.ConnectError:
        raise RuntimeError("PaddleOCR 服务不可用，请确保 OCR Docker 容器已启动。")
    except Exception as e:
        raise RuntimeError(f"调用 OCR 服务失败: {e}")

    if response.status_code != 200:
        raise RuntimeError(f"OCR 服务返回错误 (HTTP {response.status_code}): {response.text}")

    data = response.json()
    text = data.get("markdown", "")
    if not text.strip():
        raise RuntimeError("OCR 服务返回了空内容，文档可能为空或无法识别")
    return text


def _parse_docx(file_path: str) -> str:
    """Extract text from Word document."""
    logger.info(f"Parsing DOCX: {file_path}")
    doc = docx.Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Check if paragraph has a heading style
        style_name = getattr(getattr(para, 'style', None), 'name', '') or ''
        if style_name.startswith("Heading"):
            level = style_name.split()[-1]
            try:
                hashes = "#" * int(level)
            except ValueError:
                hashes = "#"
            paragraphs.append(f"{hashes} {text}")
        else:
            paragraphs.append(text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def _parse_xlsx(file_path: str) -> str:
    """Extract text from Excel spreadsheet."""
    logger.info(f"Parsing XLSX: {file_path}")
    wb = openpyxl.load_workbook(file_path, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## Sheet: {sheet_name}")
        rows_text = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows_text.append(" | ".join(cells))
        parts.append("\n".join(rows_text))
    wb.close()
    return "\n\n".join(parts)


def _parse_text(file_path: str) -> str:
    """Read plain text / Markdown file."""
    logger.info(f"Reading text file: {file_path}")
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()
